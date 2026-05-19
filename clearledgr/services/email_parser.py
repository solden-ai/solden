"""
Solden Email Parser Service

Parses finance-related emails and attachments:
- Invoice extraction from PDFs (with table support via pdfplumber)
- OCR for scanned documents and images (pytesseract)
- Payment confirmation parsing
- Bank statement detection
- Vendor context extraction with fuzzy matching
"""

import re
import zipfile
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timezone
import base64
import io
import logging

from clearledgr.core.org_utils import assert_org_id
from clearledgr.core.utils import safe_float

logger = logging.getLogger(__name__)

# Optional imports for enhanced extraction
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("pytesseract not available - OCR disabled. Scanned invoices will be flagged as 'requires_ocr'. Install with: pip install pytesseract pillow")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available - table extraction disabled. Install with: pip install pdfplumber")

try:
    import pypdfium2 as pdfium
    PDFIUM_AVAILABLE = True
except ImportError:
    PDFIUM_AVAILABLE = False
    logger.warning("pypdfium2 not available - scanned PDF OCR disabled. Install with: pip install pypdfium2")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX parsing disabled. Install with: pip install python-docx")

try:
    from rapidfuzz import fuzz, process
    FUZZY_AVAILABLE = True
except ImportError:
    FUZZY_AVAILABLE = False
    logger.warning("rapidfuzz not available - fuzzy matching disabled. Install with: pip install rapidfuzz")


# Common vendor names for fuzzy matching
KNOWN_VENDORS = [
    "Amazon", "Amazon Web Services", "AWS", "Microsoft", "Google", "Apple",
    "Stripe", "PayPal", "Shopify", "Salesforce", "HubSpot", "Slack", "Zoom",
    "Adobe", "Atlassian", "Dropbox", "GitHub", "Notion", "Figma", "Canva",
    "QuickBooks", "Xero", "FreshBooks", "Wave", "Gusto", "Deel", "Remote",
    "Office Depot", "Staples", "FedEx", "UPS", "DHL", "USPS",
    "Uber", "Lyft", "Delta", "United", "American Airlines", "Southwest",
    "Hilton", "Marriott", "Airbnb", "WeWork", "Regus",
    "Verizon", "AT&T", "T-Mobile", "Comcast", "Spectrum",
    "PG&E", "ConEd", "Duke Energy", "National Grid",
    "Bank of America", "Chase", "Wells Fargo", "Citi", "Capital One",
]

GENERIC_MAILBOX_DOMAINS = {
    "gmail.com",
    "googlemail.com",
    "yahoo.com",
    "ymail.com",
    "hotmail.com",
    "outlook.com",
    "live.com",
    "icloud.com",
    "me.com",
    "mac.com",
    "aol.com",
    "protonmail.com",
    "pm.me",
    "gmx.com",
    "zoho.com",
}


class EmailParser:
    """
    Parses email content and attachments to extract financial data.
    """
    
    # Comprehensive patterns for financial emails (international support)
    AMOUNT_PATTERNS = [
        # Currency symbols with amounts
        r'(?:€|EUR)\s*([\d\s.,]+)',  # EUR format
        r'(?:\$|USD)\s*([\d\s.,]+)',  # USD format
        r'(?:£|GBP)\s*([\d\s.,]+)',  # GBP format
        r'(?:₦|NGN)\s*([\d\s.,]+)',  # Nigerian Naira
        r'(?:GH₵|GHS|¢)\s*([\d\s.,]+)',  # Ghanaian Cedi
        r'(?:\bZAR\b|\bR(?=\s*\d))\s*([\d\s.,]+)',  # South African Rand
        r'(?:KES|KSh)\s*([\d\s.,]+)',  # Kenyan Shilling
        r'(?:¥|JPY|CNY)\s*([\d\s.,]+)',  # Japanese Yen / Chinese Yuan
        r'(?:₹|INR)\s*([\d\s.,]+)',  # Indian Rupee
        r'(?:CHF)\s*([\d\s.,]+)',  # Swiss Franc
        r'(?:AUD|A\$)\s*([\d\s.,]+)',  # Australian Dollar
        r'(?:CAD|C\$)\s*([\d\s.,]+)',  # Canadian Dollar
        r'(?:SEK|kr)\s*([\d\s.,]+)',  # Swedish Krona
        r'(?:NOK)\s*([\d\s.,]+)',  # Norwegian Krone
        r'(?:DKK)\s*([\d\s.,]+)',  # Danish Krone
        r'(?:PLN|zł)\s*([\d\s.,]+)',  # Polish Zloty
        r'(?:BRL|R\$)\s*([\d\s.,]+)',  # Brazilian Real
        r'(?:MXN)\s*([\d\s.,]+)',  # Mexican Peso
        r'(?:AED)\s*([\d\s.,]+)',  # UAE Dirham
        r'(?:SAR)\s*([\d\s.,]+)',  # Saudi Riyal
        # Amount labels (more comprehensive)
        r'Total\s*(?:Amount|Due|Payable)?[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Amount\s*(?:Due|Payable)?[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Net\s*(?:Amount|Total)?[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Grand\s+Total[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Balance\s*(?:Due)?[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Subtotal[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Invoice\s+Total[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'Pay\s+This\s+Amount[:\s]+(?:€|\$|£|₦|R|KES|¥|₹)?\s*([\d\s.,]+)',
        r'(?:Total\s*(?:payment|amount|due|due\s+amount|payable)?|Grand\s+Total|Invoice\s+Total|Balance\s*(?:Due)?|Subtotal)[:\s()\-]*([\d\s.,]+)\s*(?:GHS|GH₵|USD|EUR|GBP|NGN|ZAR|KES|JPY|CNY|INR|CHF|AUD|CAD|SGD|AED|SAR|THB)\b',
    ]
    
    INVOICE_PATTERNS = [
        r'Invoice\s*(?:Number|No\.?|#)\s*[:#-]?\s*([A-Z0-9][\w\-/]+)',
        r'\b((?:INV)(?:[:\-\s#]*[A-Z0-9][\w\-]+))\b',
        r'Bill\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'Reference\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'Order\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'PO\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'Receipt\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'Transaction\s*(?:ID|Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
        r'Doc(?:ument)?\s*(?:Number|No\.?|#)?[:\s]*([A-Z0-9][\w\-/]+)',
    ]
    
    DATE_PATTERNS = [
        # ISO format
        r'(\d{4}-\d{2}-\d{2})',
        r'(\d{4}/\d{2}/\d{2})',
        # European formats (DD/MM/YYYY, DD-MM-YYYY, DD.MM.YYYY)
        r'(\d{1,2}/\d{1,2}/\d{4})',
        r'(\d{1,2}-\d{1,2}-\d{4})',
        r'(\d{1,2}\.\d{1,2}\.\d{4})',
        # US format (MM/DD/YYYY)
        r'(\d{1,2}/\d{1,2}/\d{2,4})',
        # Written dates
        r'(\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})',
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},?\s+\d{4})',
        # With labels
        r'(?:Due|Date|Invoice\s+Date|Issue\s+Date)[:\s]+(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})',
        r'(?:Due|Date|Invoice\s+Date|Issue\s+Date)[:\s]+(\d{1,2}\s+\w+\s+\d{4})',
    ]
    
    PAYMENT_REQUEST_KEYWORDS = [
        'payment request', 'please pay', 'requesting payment',
        'reimburse', 'reimbursement', 'expense report',
        'wire to', 'transfer to', 'pay to', 'contractor payment'
    ]

    INVOICE_KEYWORDS = [
        'invoice', 'bill', 'amount due', 'balance due',
        'total due', 'payable', 'payment terms', 'due date', 'invoice number'
    ]

    RECEIPT_KEYWORDS = [
        'your receipt', 'payment receipt', 'receipt from', 'receipt for',
        'thank you for your payment', 'thank you for your purchase',
        'order confirmation', 'subscription receipt', 'order receipt',
    ]

    PAYMENT_KEYWORDS = [
        'payment confirmation', 'payment confirmed', 'payment received',
        'payment processed', 'payment successful', 'payment completed',
        'invoice paid', 'paid successfully',
    ]

    REFUND_KEYWORDS = [
        'refund', 'refunded', 'refund receipt', 'refund confirmation',
        'refund processed', 'money returned',
    ]

    CREDIT_NOTE_KEYWORDS = [
        'credit note', 'credit memo', 'vendor credit',
        'credit applied', 'credit issued',
    ]

    STATEMENT_KEYWORDS = [
        'bank statement', 'card statement', 'account statement',
        'monthly statement', 'billing statement',
    ]

    # Domains that act as payment processors / billing platforms.
    # When the sender is one of these, the actual vendor is in the email body.
    PAYMENT_PROCESSOR_DOMAINS = {
        'stripe.com', 'paypal.com', 'square.com', 'squareup.com',
        'braintree.com', 'paddle.com', 'chargebee.com', 'recurly.com',
        'fastspring.com', 'gumroad.com', 'lemonsqueezy.com',
        'bill.com', 'payoneer.com', 'wise.com', 'transferwise.com',
    }
    
    def __init__(self):
        self.supported_currencies = [
            'EUR', 'USD', 'GBP', 'NGN', 'ZAR', 'KES',
            'GHS', 'JPY', 'CNY', 'INR', 'CHF', 'AUD', 'CAD',
            'SEK', 'NOK', 'DKK', 'PLN', 'BRL', 'MXN',
            'AED', 'SAR', 'SGD', 'HKD', 'NZD', 'THB',
        ]
        self.known_vendors = KNOWN_VENDORS.copy()
    
    def parse_email(
        self,
        subject: str,
        body: str,
        sender: str,
        attachments: List[Dict] = None
    ) -> Dict[str, Any]:
        """
        Parse an email and extract financial data.
        
        Args:
            subject: Email subject line
            body: Email body text
            sender: Sender email address
            attachments: List of attachments [{name, content_type, content_base64}]
            
        Returns:
            Parsed email data with extracted fields
        """
        attachments = attachments or []

        # C12: Detect forwarded emails and extract from the inner (forwarded) content
        forwarded = False
        forwarded_markers = [
            "---------- Forwarded message ---------",
            "Begin forwarded message",
        ]
        for marker in forwarded_markers:
            marker_pos = body.find(marker)
            if marker_pos != -1:
                forwarded = True
                # Extract the inner forwarded content for vendor/amount extraction
                body = body[marker_pos + len(marker):]
                break

        # Determine email type
        email_type = self._classify_email(subject, body)

        # Extract vendor from sender (passes subject+body for payment-processor senders)
        vendor = self._extract_vendor(sender, subject=subject, body=body)

        # Extract amounts
        amounts = self._extract_amounts(subject + " " + body)
        
        # Extract invoice numbers
        invoice_numbers = self._extract_invoice_numbers(subject + " " + body)
        
        # Extract dates
        dates = self._extract_dates(subject + " " + body)

        email_fields = {
            "vendor": vendor,
            "amount": self._primary_amount_details(amounts).get("value"),
            "currency": self._primary_amount_details(amounts).get("currency"),
            "invoice_number": invoice_numbers[0] if invoice_numbers else None,
            "invoice_date": dates[0] if dates else None,
            "due_date": None,
            "amount_score": self._primary_amount_details(amounts).get("score"),
        }
        
        # Parse attachments
        parsed_attachments = []
        attachment_extractions = []
        primary_source = "email"
        preferred_extraction: Dict[str, Any] = {}
        attachment_fields = {
            "vendor": None,
            "amount": None,
            "currency": None,
            "invoice_number": None,
            "invoice_date": None,
            "due_date": None,
            "amount_score": 1.0,
        }
        for attachment in attachments:
            att_name = (attachment.get('name') or attachment.get('filename') or '').lower()
            att_content_base64 = attachment.get('content_base64')

            # ZIP archive: extract inner files and process each individually
            if att_name.endswith('.zip') and att_content_base64:
                archive_files = self._extract_archive_attachments(
                    att_content_base64,
                    attachment.get('name') or attachment.get('filename') or 'unknown.zip',
                )
                for af in archive_files:
                    result = self._parse_attachment(af)
                    if result:
                        parsed_attachments.append(result)
                        extraction = result.get("extraction")
                        if isinstance(extraction, dict) and extraction:
                            attachment_extractions.append(extraction)
                continue  # Skip processing the ZIP itself

            # RAR: not supported (requires third-party library)
            if att_name.endswith('.rar'):
                logger.warning("RAR archive detected (%s) — skipping (not supported)", att_name)
                continue

            parsed = self._parse_attachment(attachment)
            if parsed:
                parsed_attachments.append(parsed)
                extraction = parsed.get("extraction")
                if isinstance(extraction, dict) and extraction:
                    attachment_extractions.append(extraction)

            attachment_text = attachment.get("content_text")
            if attachment_text:
                parsed_text = self.parse_invoice_text(attachment_text)
                if parsed_text:
                    attachment_extractions.append(parsed_text)
                    attachment_name = attachment.get("name") or attachment.get("filename")
                    attachment_type = attachment.get("content_type") or attachment.get("mime_type")
                    parsed_attachments.append({
                        "name": attachment_name,
                        "type": parsed_text.get("type") or (parsed.get("type") if parsed else "document"),
                        "content_type": attachment_type,
                        "parsed": True,
                        "extraction": parsed_text
                    })

        if attachment_extractions:
            primary_source = "attachment"
            preferred_extraction = next(
                (
                    extraction for extraction in attachment_extractions
                    if str(extraction.get("type") or "").strip().lower() == "invoice"
                ),
                attachment_extractions[0],
            )
            parsed_amount = preferred_extraction.get("amount")
            if parsed_amount:
                if isinstance(parsed_amount, dict):
                    amounts = [parsed_amount]
                else:
                    amount_value = self._parse_amount_value(
                        str(parsed_amount),
                        source_fragment=str(parsed_amount),
                    )
                    if amount_value is not None:
                        attachment_fields["amount"] = amount_value
                        attachment_fields["currency"] = preferred_extraction.get("currency") or self._detect_currency(str(parsed_amount))
                        amounts = [{
                            "value": amount_value,
                            "raw": str(parsed_amount),
                            "currency": preferred_extraction.get("currency") or self._detect_currency(str(parsed_amount)),
                        }]
            if preferred_extraction.get("invoice_number"):
                attachment_fields["invoice_number"] = preferred_extraction.get("invoice_number")
                invoice_numbers = [preferred_extraction.get("invoice_number")]
            if preferred_extraction.get("date"):
                attachment_fields["invoice_date"] = preferred_extraction.get("date")
                dates = [preferred_extraction.get("date")]
            if preferred_extraction.get("due_date"):
                attachment_fields["due_date"] = preferred_extraction.get("due_date")
                dates = [preferred_extraction.get("due_date"), *dates]
            if preferred_extraction.get("vendor"):
                attachment_fields["vendor"] = preferred_extraction.get("vendor")
                vendor = preferred_extraction.get("vendor")

        # Merge attachment data with email data
        if parsed_attachments:
            # Use attachment data if more complete
            for att in parsed_attachments:
                if att.get('amounts') and not amounts:
                    amounts = att['amounts']
                if att.get('invoice_numbers') and not invoice_numbers:
                    invoice_numbers = att['invoice_numbers']

        if invoice_numbers and amounts:
            amounts = self._filter_amounts_against_invoice_numbers(amounts, invoice_numbers)

        # Payment requests often lack invoice labels. If no amount survived the
        # main invoice-oriented extraction heuristics, fall back to a broader
        # currency+amount scan on the email text.
        if email_type == "payment_request" and not amounts:
            amounts = self._extract_payment_request_amounts(subject + " " + body)

        primary_amount = None
        primary_currency = None
        if amounts:
            if isinstance(amounts[0], dict):
                primary_amount = amounts[0].get("value")
                primary_currency = amounts[0].get("currency")
            else:
                primary_amount = amounts[0]

        if attachment_fields["amount"] is None:
            attachment_amount = preferred_extraction.get("amount")
            if isinstance(attachment_amount, dict):
                attachment_fields["amount"] = safe_float(attachment_amount.get("value"))
                attachment_fields["currency"] = attachment_fields["currency"] or attachment_amount.get("currency")
            elif attachment_amount is not None:
                attachment_fields["amount"] = self._parse_amount_value(str(attachment_amount), source_fragment=str(attachment_amount))
        attachment_fields["vendor"] = attachment_fields["vendor"] or preferred_extraction.get("vendor")
        attachment_fields["invoice_number"] = attachment_fields["invoice_number"] or preferred_extraction.get("invoice_number")
        attachment_fields["invoice_date"] = attachment_fields["invoice_date"] or preferred_extraction.get("date")
        attachment_fields["due_date"] = attachment_fields["due_date"] or preferred_extraction.get("due_date")

        final_fields = {
            "vendor": vendor,
            "amount": primary_amount,
            "currency": primary_currency,
            "invoice_number": invoice_numbers[0] if invoice_numbers else None,
            "invoice_date": self._first_matching_date(dates, attachment_fields.get("invoice_date")),
            "due_date": self._first_matching_date(dates, attachment_fields.get("due_date")),
        }
        field_provenance = self._build_field_provenance(
            final_fields=final_fields,
            email_fields=email_fields,
            attachment_fields=attachment_fields,
        )
        field_evidence = self._build_field_evidence(
            field_provenance=field_provenance,
            email_fields=email_fields,
            attachment_fields=attachment_fields,
            parsed_attachments=parsed_attachments,
        )
        source_conflicts = self._build_source_conflicts(
            email_fields=email_fields,
            attachment_fields=attachment_fields,
            field_provenance=field_provenance,
        )
        conflict_actions = self._build_conflict_actions(source_conflicts)

        # Propagate ocr_status from any attachment that needs OCR but couldn't get it
        attachments_requiring_ocr = [
            a for a in parsed_attachments
            if a.get("ocr_status") == "requires_ocr"
        ]
        ocr_status = "requires_ocr" if attachments_requiring_ocr else None

        # --- Multi-invoice detection ---
        # When multiple parsed attachments each contain a distinct invoice
        # (different vendor OR invoice number OR amount), expose them as
        # separate extraction results so the triage layer can create one AP
        # item per invoice.
        distinct_invoices = self._detect_distinct_invoices(
            attachment_extractions, parsed_attachments, email_fields
        )
        multiple_invoices = len(distinct_invoices) > 1

        result = {
            "email_type": email_type,
            "vendor": vendor,
            "sender": sender,
            "subject": subject,
            "amounts": amounts,
            "primary_amount": primary_amount,
            "invoice_numbers": invoice_numbers,
            "primary_invoice": invoice_numbers[0] if invoice_numbers else None,
            "dates": dates,
            "primary_date": dates[0] if dates else None,
            "attachments": parsed_attachments,
            "has_invoice_attachment": any(a.get('type') == 'invoice' for a in parsed_attachments),
            "has_statement_attachment": any(a.get('type') == 'statement' for a in parsed_attachments),
            "confidence": self._calculate_confidence(email_type, amounts, invoice_numbers),
            "currency": primary_currency,
            "primary_source": primary_source,
            "ocr_status": ocr_status,
            "field_provenance": field_provenance,
            "field_evidence": field_evidence,
            "source_conflicts": source_conflicts,
            "requires_extraction_review": any(bool(entry.get("blocking")) for entry in source_conflicts),
            "conflict_actions": conflict_actions,
            "forwarded": forwarded,
            "attachment_count": len(attachments),
            "invoice_count": len(distinct_invoices) if multiple_invoices else 1,
            "multiple_invoices": multiple_invoices,
            "parsed_at": datetime.now(timezone.utc).isoformat()
        }
        if multiple_invoices:
            result["invoices"] = distinct_invoices
        return result

    def _detect_distinct_invoices(
        self,
        attachment_extractions: List[Dict[str, Any]],
        parsed_attachments: List[Dict[str, Any]],
        email_fields: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """Return a list of distinct invoice extraction dicts when multiple
        attachments each represent a separate invoice.

        An invoice extraction is considered *distinct* when at least one of
        (vendor, invoice_number, amount) differs from every other extraction.
        Extractions that are not invoice-typed are skipped.
        """
        invoice_extractions: List[Dict[str, Any]] = []
        for idx, att in enumerate(parsed_attachments):
            ext = att.get("extraction")
            if not isinstance(ext, dict) or not ext:
                continue
            att_type = str(att.get("type") or "").strip().lower()
            if att_type not in ("invoice", "document", ""):
                continue
            # Need at least one identifying field
            vendor = str(ext.get("vendor") or "").strip().lower()
            inv_num = str(ext.get("invoice_number") or "").strip().lower()
            amount = safe_float(ext.get("amount"))
            if not vendor and not inv_num and amount == 0.0:
                continue
            invoice_extractions.append({
                "vendor": ext.get("vendor") or email_fields.get("vendor"),
                "amount": safe_float(ext.get("amount")),
                "currency": ext.get("currency") or email_fields.get("currency"),
                "invoice_number": ext.get("invoice_number"),
                "invoice_date": ext.get("date") or ext.get("invoice_date"),
                "due_date": ext.get("due_date"),
                "confidence": safe_float(ext.get("confidence"), 0.5),
                "attachment_index": idx,
                "attachment_name": att.get("name"),
                "_key": (vendor, inv_num, amount),
            })

        if len(invoice_extractions) < 2:
            return invoice_extractions

        # Deduplicate: keep only truly distinct invoices
        distinct: List[Dict[str, Any]] = []
        seen_keys: set = set()
        for inv in invoice_extractions:
            key = inv.pop("_key")
            if key not in seen_keys:
                seen_keys.add(key)
                distinct.append(inv)
        return distinct

    def _primary_amount_details(self, amounts: List[Any]) -> Dict[str, Any]:
        if not isinstance(amounts, list) or not amounts:
            return {"value": None, "currency": None, "score": 0.0, "raw": None}
        primary = amounts[0]
        if isinstance(primary, dict):
            return {
                "value": primary.get("value"),
                "currency": primary.get("currency"),
                "score": safe_float(primary.get("score"), 0.0),
                "raw": primary.get("raw"),
            }
        return {"value": primary, "currency": None, "score": 0.0, "raw": primary}

    def _has_field_value(self, value: Any) -> bool:
        if value is None:
            return False
        if isinstance(value, str):
            return bool(value.strip())
        return True

    def _first_matching_date(self, dates: List[Any], preferred: Any) -> Optional[str]:
        preferred_token = str(preferred or "").strip()
        if preferred_token:
            return preferred_token
        for value in dates or []:
            token = str(value or "").strip()
            if token:
                return token
        return None

    def _generic_vendor_candidate(self, value: Any) -> bool:
        token = re.sub(r"\s+", " ", str(value or "").strip()).lower()
        if not token:
            return True
        if token in {"unknown", "unknown vendor", "vendor", "merchant", "payment processor"}:
            return True
        return any(marker in token for marker in ("payment", "payments", "billing", "noreply"))

    def _comparable_field_value(self, field: str, value: Any) -> Any:
        if value is None:
            return None
        if field == "amount":
            try:
                return round(float(value), 2)
            except (TypeError, ValueError):
                return None
        if field == "currency":
            token = str(value or "").strip().upper()
            return token or None
        token = str(value or "").strip()
        if not token:
            return None
        if field == "vendor":
            return re.sub(r"[^a-z0-9]", "", token.lower()) or None
        return re.sub(r"\s+", "", token).lower()

    def _sources_conflict(self, field: str, left: Any, right: Any) -> bool:
        left_value = self._comparable_field_value(field, left)
        right_value = self._comparable_field_value(field, right)
        if left_value is None or right_value is None:
            return False
        if field == "vendor" and (
            self._generic_vendor_candidate(left) or self._generic_vendor_candidate(right)
        ):
            return False
        return left_value != right_value

    def _build_field_provenance(
        self,
        *,
        final_fields: Dict[str, Any],
        email_fields: Dict[str, Any],
        attachment_fields: Dict[str, Any],
    ) -> Dict[str, Dict[str, Any]]:
        provenance: Dict[str, Dict[str, Any]] = {}
        for field in ("vendor", "amount", "currency", "invoice_number", "invoice_date", "due_date"):
            final_value = final_fields.get(field)
            email_value = email_fields.get(field)
            attachment_value = attachment_fields.get(field)
            chosen_source = None
            if self._has_field_value(final_value):
                if self._has_field_value(attachment_value) and not self._sources_conflict(field, final_value, attachment_value):
                    chosen_source = "attachment"
                elif self._has_field_value(email_value) and not self._sources_conflict(field, final_value, email_value):
                    chosen_source = "email"
            if not chosen_source:
                chosen_source = "attachment" if self._has_field_value(attachment_value) else "email"
            candidates = {}
            if self._has_field_value(email_value):
                candidates["email"] = email_value
            if self._has_field_value(attachment_value):
                candidates["attachment"] = attachment_value
            provenance[field] = {
                "source": chosen_source,
                "value": final_value,
                "candidates": candidates,
            }
        return provenance

    def _build_field_evidence(
        self,
        *,
        field_provenance: Dict[str, Dict[str, Any]],
        email_fields: Dict[str, Any],
        attachment_fields: Dict[str, Any],
        parsed_attachments: List[Dict[str, Any]],
    ) -> Dict[str, Dict[str, Any]]:
        attachment_name = None
        for parsed in parsed_attachments:
            token = str(parsed.get("name") or "").strip()
            if token:
                attachment_name = token
                break
        evidence: Dict[str, Dict[str, Any]] = {}
        for field, entry in field_provenance.items():
            source = str(entry.get("source") or "email").strip() or "email"
            field_evidence = {
                "source": source,
                "selected_value": entry.get("value"),
                "email_value": email_fields.get(field),
                "attachment_value": attachment_fields.get(field),
            }
            if source == "attachment" and attachment_name:
                field_evidence["attachment_name"] = attachment_name
            evidence[field] = {
                key: value
                for key, value in field_evidence.items()
                if value not in (None, "", [], {})
            }
        return evidence

    def _build_source_conflicts(
        self,
        *,
        email_fields: Dict[str, Any],
        attachment_fields: Dict[str, Any],
        field_provenance: Dict[str, Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        conflicts: List[Dict[str, Any]] = []
        for field in ("amount", "currency", "invoice_number", "vendor", "due_date"):
            email_value = email_fields.get(field)
            attachment_value = attachment_fields.get(field)
            if not self._sources_conflict(field, email_value, attachment_value):
                continue
            blocking = field in {"amount", "currency", "invoice_number"}
            conflicts.append(
                {
                    "field": field,
                    "reason": "source_value_mismatch",
                    "severity": "high" if blocking else "medium",
                    "blocking": blocking,
                    "preferred_source": field_provenance.get(field, {}).get("source") or "attachment",
                    "values": {
                        "email": email_value,
                        "attachment": attachment_value,
                    },
                }
            )
        return conflicts

    def _build_conflict_actions(self, source_conflicts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        actions: List[Dict[str, Any]] = []
        for conflict in source_conflicts:
            field = str(conflict.get("field") or "").strip()
            if not field:
                continue
            actions.append(
                {
                    "action": "review_fields",
                    "field": field,
                    "reason": str(conflict.get("reason") or "source_value_mismatch"),
                    "blocking": bool(conflict.get("blocking")),
                }
            )
        return actions
    
    def parse_invoice_text(self, text: str) -> Dict[str, Any]:
        """
        Parse invoice text (from PDF extraction or OCR).
        
        Args:
            text: Extracted text from invoice
            
        Returns:
            Parsed invoice data
        """
        amounts = self._extract_amounts(text)
        invoice_numbers = self._extract_invoice_numbers(text)
        dates = self._extract_dates(text)
        
        # Try to extract line items
        line_items = self._extract_line_items(text)
        
        # Extract vendor name (usually at top of invoice)
        vendor = self._extract_vendor_from_text(text)
        
        return {
            "type": "invoice",
            "vendor": vendor,
            "invoice_number": invoice_numbers[0] if invoice_numbers else None,
            "amount": amounts[0] if amounts else None,
            "all_amounts": amounts,
            "date": dates[0] if dates else None,
            "due_date": self._extract_due_date(text),
            "line_items": line_items,
            "currency": self._detect_currency(text),
            "parsed_at": datetime.now(timezone.utc).isoformat()
        }
    
    def parse_payment_confirmation(self, text: str) -> Dict[str, Any]:
        """
        Parse payment confirmation email.
        
        Args:
            text: Email body text
            
        Returns:
            Parsed payment data
        """
        amounts = self._extract_amounts(text)
        
        # Extract transaction ID
        txn_patterns = [
            r'Transaction\s*(?:ID|#|Number)?[:\s]+([A-Z0-9\-_]+)',
            r'Reference[:\s]+([A-Z0-9\-_]+)',
            r'Confirmation[:\s]+([A-Z0-9\-_]+)',
            r'TXN[:\-\s]*([A-Z0-9\-]+)',
        ]
        
        txn_id = None
        for pattern in txn_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                txn_id = match.group(1)
                break
        
        # Extract payer/payee
        payer = self._extract_party(text, 'from')
        payee = self._extract_party(text, 'to')
        
        return {
            "type": "payment",
            "transaction_id": txn_id,
            "amount": amounts[0] if amounts else None,
            "currency": self._detect_currency(text),
            "payer": payer,
            "payee": payee,
            "date": self._extract_dates(text)[0] if self._extract_dates(text) else None,
            "status": "completed",
            "parsed_at": datetime.now(timezone.utc).isoformat()
        }
    
    def _classify_email(self, subject: str, body: str) -> str:
        """Classify email type based on content.

        Credit notes and refunds must win before receipt/invoice checks because
        they often mention invoice IDs and payment receipts in the body.
        """
        text = (subject + " " + body).lower()

        if any(kw in text for kw in self.CREDIT_NOTE_KEYWORDS):
            return "credit_note"

        if any(kw in text for kw in self.REFUND_KEYWORDS):
            return "refund"

        if any(kw in text for kw in self.STATEMENT_KEYWORDS):
            return "statement"

        if any(kw in text for kw in self.PAYMENT_KEYWORDS):
            return "payment"

        # Receipts must be detected before invoice keywords to avoid misclassification.
        if any(kw in text for kw in self.RECEIPT_KEYWORDS):
            return "receipt"

        if any(kw in text for kw in self.INVOICE_KEYWORDS):
            return "invoice"

        if any(kw in text for kw in self.PAYMENT_REQUEST_KEYWORDS):
            return "payment_request"

        return "general"

    def _extract_vendor_from_email_context(self, subject: str, body: str) -> Optional[str]:
        """Extract actual vendor name from subject/body when the sender is a payment processor.

        Handles common patterns:
          - "Your receipt from Replit #2462-2703"  → "Replit"
          - "Invoice from Acme Corp"                → "Acme Corp"
        """
        import re
        lines = [re.sub(r"\s+", " ", str(line or "")).strip() for line in str(body or "").splitlines()]
        label_candidates = self._extract_vendor_from_labeled_lines(
            lines,
            labels={"from", "issued by", "payee", "vendor", "supplier", "seller", "merchant"},
        )
        if label_candidates:
            return label_candidates

        patterns = [
            # "receipt from X #123" / "payment confirmation from X" / "credit note from Acme Corp"
            r'(?:receipt|payment(?:\s+confirmation)?|refund|invoice|bill|statement|credit\s+note|credit\s+memo)\s+from\s+([A-Z][A-Za-z0-9\s&\.\-]+?)(?:\s+#|\s+\d|\s+for\b|[,\.]|$)',
            # "Your receipt from X" / "Your payment confirmation from X"
            r'^(?:your\s+)?(?:receipt|payment(?:\s+confirmation)?|refund|invoice|bill|statement|order|credit\s+note|credit\s+memo)\s+from\s+([A-Z][A-Za-z0-9\s&\.\-]+?)(?:\s+#|\s+\d|[,\.]|$)',
            # "Acme invoice", "Acme receipt", "Acme payment confirmation"
            r'^([A-Z][A-Za-z0-9\s&\.\-]+?)\s+(?:invoice|receipt|payment(?:\s+confirmation)?|bill|refund|credit\s+note|credit\s+memo)\b',
            # "Tuition fees from X", "Charges from X"
            r'(?:fees|charges|tuition\s+fees|statement|invoice|bill)\s+from\s+([A-Z][A-Za-z0-9\s&\.\-]+?)(?:\s+#|\s+\d|\s+for\b|[,\.]|$)',
        ]
        for text in [subject, (body or '')[:300]]:
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    candidate = match.group(1).strip().rstrip('.')
                    if 2 < len(candidate) < 50 and candidate.lower() not in {
                        'your', 'the', 'a', 'an', 'this', 'our', 'my'
                    }:
                        return self._normalize_vendor_candidate(candidate)
        return None

    def _extract_vendor(self, sender: str, subject: str = '', body: str = '') -> str:
        """Extract vendor name from sender email with fuzzy matching.

        For known payment-processor domains (Stripe, PayPal, etc.) the real
        vendor is the merchant, not the processor.  Fall back to extracting
        the vendor name from the subject/body in those cases.
        """
        sender_text = str(sender or "").strip()
        email_match = re.search(r'<([^>]+@[^>]+)>', sender_text) or re.search(r'([A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,})', sender_text, re.IGNORECASE)
        sender_email = email_match.group(1) if email_match else sender_text
        if '@' in sender_email:
            domain = sender_email.split('@')[1].lower().strip(" >")
            # Strip subdomains to get the base domain (e.g. "invoice.stripe.com" → "stripe.com")
            parts = domain.rsplit('.', 2)
            base_domain = '.'.join(parts[-2:]) if len(parts) >= 2 else domain

            if base_domain in self.PAYMENT_PROCESSOR_DOMAINS or base_domain in GENERIC_MAILBOX_DOMAINS:
                extracted = self._extract_vendor_from_email_context(subject, body)
                if extracted:
                    return extracted
                # Fall through to domain-based extraction as last resort

            name = domain.split('.')[0]
            capitalized = name.title()
            return self._normalize_vendor_candidate(capitalized)
        return sender

    def _normalize_vendor_candidate(self, candidate: Optional[str]) -> Optional[str]:
        """Normalize vendor names against a known-vendor list conservatively.

        We prefer avoiding false positives (e.g., ``Taskforce`` -> ``Salesforce``)
        over aggressively normalizing unknown vendors.
        """
        if not candidate:
            return candidate

        cleaned = str(candidate).strip()
        cleaned = re.sub(
            r"\s+(?:payer|bill\s+to|billed\s+to|recipient|customer)\b.*$",
            "",
            cleaned,
            flags=re.IGNORECASE,
        ).strip(" .,:;|-")
        if not cleaned:
            return cleaned

        if not (FUZZY_AVAILABLE and self.known_vendors):
            return cleaned

        def _norm(text: str) -> str:
            return "".join(ch.lower() for ch in str(text) if ch.isalnum())

        candidate_norm = _norm(cleaned)
        if not candidate_norm:
            return cleaned

        normalized_map: Dict[str, str] = {}
        for vendor in self.known_vendors:
            normalized_map.setdefault(_norm(vendor), vendor)

        # Exact normalized match handles casing and punctuation variants safely
        # (e.g., Aws -> AWS, Hubspot -> HubSpot).
        exact = normalized_map.get(candidate_norm)
        if exact:
            return exact

        match = process.extractOne(
            candidate_norm,
            list(normalized_map.keys()),
            scorer=fuzz.ratio,
            score_cutoff=75,
        )
        if not match:
            return cleaned

        matched_norm = match[0]
        score = float(match[1])

        # Require high confidence, or moderate confidence with strong shared prefix.
        shared_prefix = 0
        for left, right in zip(candidate_norm, matched_norm):
            if left != right:
                break
            shared_prefix += 1

        if score >= 88 or (score >= 75 and shared_prefix >= 4):
            return normalized_map.get(matched_norm, cleaned)

        return cleaned
    
    def _extract_vendor_from_text(self, text: str) -> Optional[str]:
        """Extract vendor name from invoice text with fuzzy matching."""
        lines = [re.sub(r"\s+", " ", str(line or "")).strip(" .,:;|-") for line in text.split('\n')]
        labeled = self._extract_vendor_from_labeled_lines(
            lines,
            labels={"issued by", "payee", "vendor", "supplier", "seller", "merchant"},
        )
        if labeled:
            return labeled

        candidates = []
        
        # Look for company name patterns at start of text
        lines = text.split('\n')[:40]  # Early document lines carry vendor identity
        
        for line in lines:
            line = re.sub(r"\.(?=\w)", "", line)
            line = re.sub(r"\s+", " ", line).strip(" .,:;|-")
            if line.startswith("--- Page"):
                continue
            if self._is_vendor_noise_candidate(line):
                continue
            if line and line[0].isupper():
                candidates.append(line)
        
        if not candidates:
            return None

        legal_entity_pattern = re.compile(
            r"\b(?:limited|ltd|llc|inc|corp|corporation|gmbh|plc|pte|pty|bv|sarl|sa)\b",
            re.IGNORECASE,
        )
        for candidate in candidates:
            if legal_entity_pattern.search(candidate):
                return candidate

        for candidate in candidates:
            if len(candidate.split()) >= 2 and len(candidate) >= 8:
                return candidate
        
        # Try fuzzy match against known vendors
        if FUZZY_AVAILABLE and self.known_vendors:
            for candidate in candidates:
                match = process.extractOne(
                    candidate,
                    self.known_vendors,
                    scorer=fuzz.partial_ratio,
                    score_cutoff=75
                )
                if match:
                    return match[0]  # Return standardized vendor name

        # Return first candidate if no fuzzy match
        return candidates[0] if candidates else None

    def _extract_vendor_from_labeled_lines(
        self,
        lines: List[str],
        *,
        labels: set[str],
    ) -> Optional[str]:
        normalized_lines = [re.sub(r"\s+", " ", str(line or "")).strip(" .,:;|-") for line in lines or []]
        for index, line in enumerate(normalized_lines[:120]):
            lowered = line.lower().rstrip(":")
            inline_match = re.match(
                r"^(from|issued by|payee|vendor|supplier|seller|merchant)\s*:\s*(.+)$",
                line,
                re.IGNORECASE,
            )
            if inline_match:
                label = inline_match.group(1).strip().lower()
                if label == "from":
                    value = re.sub(r"\s*<[^>]+>$", "", inline_match.group(2)).strip()
                else:
                    value = inline_match.group(2).strip()
                candidate = self._normalize_vendor_candidate(value)
                if label in labels and candidate and not self._is_vendor_noise_candidate(candidate):
                    return candidate
            if lowered not in labels:
                continue
            for next_line in normalized_lines[index + 1:index + 4]:
                if not next_line:
                    continue
                candidate = re.sub(r"\s*<[^>]+>$", "", next_line).strip()
                candidate = self._normalize_vendor_candidate(candidate)
                if candidate and not self._is_vendor_noise_candidate(candidate):
                    return candidate
        return None

    def _is_vendor_noise_candidate(self, candidate: Optional[str]) -> bool:
        cleaned = re.sub(r"\s+", " ", str(candidate or "")).strip(" .,:;|-")
        if len(cleaned) < 3 or len(cleaned) > 80:
            return True
        lowered = cleaned.lower()
        normalized_letters = re.sub(r"[^a-z]", "", lowered)
        if re.search(r"\d{4}", cleaned):
            return True
        if any(
            token in lowered or token.replace(" ", "") in normalized_letters
            for token in (
                "invoice",
                "bill to",
                "details",
                "summary",
                "subtotal",
                "total",
                "vat",
                "billing id",
                "account id",
                "description",
                "amount",
                "due",
                "page ",
            )
        ):
            return True
        if cleaned.count(".") >= 3:
            return True
        if re.search(r"\b(?:usd|eur|gbp|cad|aud|inr|ngn|kes|jpy|cny|vat)\b", lowered):
            return True
        if any(symbol in cleaned for symbol in ("$", "€", "£", "¥", "₹", "₦")):
            return True
        if lowered in {"pay online", "united states"}:
            return True
        if any(ch.isdigit() for ch in cleaned):
            return True
        alnum_chars = [ch for ch in cleaned if ch.isalnum()]
        if alnum_chars:
            alpha_ratio = sum(ch.isalpha() for ch in alnum_chars) / len(alnum_chars)
            if alpha_ratio < 0.75:
                return True
        compact = re.sub(r"[^A-Za-z]", "", cleaned)
        if compact and len(compact) <= 4:
            return True
        tokens = [token for token in cleaned.split() if token]
        if len(tokens) >= 3 and sum(len(token) <= 2 for token in tokens) >= 2:
            return True
        return False
    
    def add_known_vendor(self, vendor_name: str):
        """Add a vendor to the known vendors list for fuzzy matching."""
        if vendor_name and vendor_name not in self.known_vendors:
            self.known_vendors.append(vendor_name)
    
    def _extract_amounts(self, text: str) -> List[Dict[str, Any]]:
        """Extract monetary amounts from text."""
        amounts = []

        for pattern in self.AMOUNT_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                raw = match.group(1) if match.groups() else match.group(0)
                value = self._parse_amount_value(raw, source_fragment=match.group(0))
                # Keep legitimate zero-value invoices (e.g., $0.00 credit/settled cycles).
                if value is None or value < 0:
                    continue

                context = text[max(0, match.start() - 40):match.end() + 40].lower()
                if self._looks_like_identifier_token(raw):
                    continue
                score = self._score_amount_candidate(match.group(0), context)
                local_currency = self._detect_currency(f"{match.group(0)} {context}") or self._detect_currency(text)

                amounts.append({
                    "value": value,
                    "raw": raw,
                    "currency": local_currency,
                    "score": score
                })

        # Remove duplicates, keep the highest-scoring candidate for each value
        unique_map: Dict[float, Dict[str, Any]] = {}
        for a in amounts:
            value = a["value"]
            if value not in unique_map or a.get("score", 0) > unique_map[value].get("score", 0):
                unique_map[value] = a

        unique = list(unique_map.values())
        # Prefer labeled totals, then larger values
        return sorted(unique, key=lambda x: (x.get("score", 0), x["value"]), reverse=True)

    def _parse_amount_value(
        self,
        raw: str,
        source_fragment: Optional[str] = None,
    ) -> Optional[float]:
        """
        Parse amount value with international format support.
        Handles: 1,234.56 | 1.234,56 | 1 234,56 | 1234.56
        """
        if raw is None:
            return None
        
        # Remove currency symbols and whitespace
        cleaned = str(raw).strip()
        cleaned = re.sub(r'[€$£₦₹¥฿]', '', cleaned)
        cleaned = cleaned.replace(' ', '').replace('\u00a0', '')  # Remove nbsp
        cleaned = cleaned.replace("'", "")
        # Amount patterns can capture trailing punctuation (e.g., "40.23.").
        # Trim non-numeric boundary characters so decimal parsing is stable.
        cleaned = cleaned.strip(".,;:-_()[]{}")
        if cleaned.endswith("-"):
            cleaned = f"-{cleaned[:-1]}"
        
        if not cleaned:
            return None

        has_comma = ',' in cleaned
        has_dot = '.' in cleaned
        normalized = cleaned

        if has_comma and has_dot:
            # Determine format based on position
            # 1,234.56 vs 1.234,56
            if cleaned.rfind(',') > cleaned.rfind('.'):
                # European: 1.234,56 -> 1234.56
                normalized = cleaned.replace('.', '').replace(',', '.')
            else:
                # US: 1,234.56 -> 1234.56
                normalized = cleaned.replace(',', '')
        elif has_comma and not has_dot:
            parts = cleaned.split(',')
            # Check if this is decimal (1234,56) or thousand separator (1,234,567)
            if len(parts) == 2 and len(parts[1]) <= 2:
                # Decimal: 1234,56 -> 1234.56
                normalized = parts[0] + '.' + parts[1]
            else:
                # Thousand separator: 1,234,567 -> 1234567
                normalized = cleaned.replace(',', '')
        elif has_dot:
            # Could be decimal or thousand separator
            parts = cleaned.split('.')
            if len(parts) == 2 and len(parts[1]) <= 2:
                # Decimal: 1234.56 (keep as is)
                normalized = cleaned
            elif len(parts) > 2:
                # Thousand separator: 1.234.567 -> 1234567
                normalized = cleaned.replace('.', '')
            else:
                normalized = cleaned
        else:
            normalized = cleaned

        try:
            value = float(normalized)
            # Validate amount is reasonable (not negative, not astronomical)
            if value < 0:
                return None
            if value > 100000000:  # 100 million cap
                return None

            # Filter obvious years (e.g., 2024, 2025) unless currency symbols present.
            raw_str = str(source_fragment or raw)
            has_currency = bool(re.search(r"(USD|EUR|GBP|\$|€|£)", raw_str, re.IGNORECASE))
            if not has_currency and value.is_integer() and 1900 <= value <= 2100:
                return None

            return value
        except ValueError:
            return None

    def _score_amount_candidate(self, fragment: str, context: str) -> int:
        """Score amount candidates so invoice totals outrank line-item values."""
        score = 0
        fragment_lower = str(fragment or "").lower()
        context_lower = str(context or "").lower()

        # Strong fragment-level signals (most precise; keep these weighted high).
        if re.search(r"(grand\s+total|total\s+due|balance\s+due|amount\s+due|invoice\s+total|amount\s+payable)", fragment_lower):
            score += 8
        elif re.search(r"\btotal\b", fragment_lower):
            score += 4

        # Penalize line-item and non-final totals aggressively so they do not
        # outrank final payable amounts when both appear in the same window.
        if re.search(r"\bsubtotal\b", fragment_lower):
            score -= 6
        if re.search(r"\b(tax|vat|gst|discount|shipping|fee|unit\s+price|qty|quantity)\b", fragment_lower):
            score -= 4

        # Context-level signals are weaker than fragment labels.
        if re.search(
            r"(total\s+due|balance\s+due|amount\s+due|invoice\s+total|grand\s+total|amount\s+payable|pay\s+this\s+amount)",
            context_lower,
        ):
            score += 2
        elif re.search(r"\btotal\b", context_lower):
            score += 1

        if re.search(r"\b(subtotal|tax|vat|gst|discount|shipping|fee|unit\s+price|qty|quantity)\b", context_lower):
            score -= 1

        # Payment request contexts often have only one amount and little invoice structure.
        if re.search(r"(payment\s+request|please\s+pay|reimburse|reimbursement)", context_lower):
            score += 2

        if re.search(r"(usd|eur|gbp|\$|€|£|cad|aud|inr|ngn)", fragment_lower, re.IGNORECASE):
            score += 1
        return score

    def _extract_payment_request_amounts(self, text: str) -> List[Dict[str, Any]]:
        """Broad currency+amount fallback for payment-request emails.

        This intentionally favors capture over invoice-specific labeling when
        `_extract_amounts()` finds nothing, while still returning scored
        candidates with currency hints.
        """
        if not text:
            return []

        candidates: List[Dict[str, Any]] = []
        patterns = [
            r"((?:USD|EUR|GBP|CAD|AUD|INR|NGN|KES|JPY|CNY|CHF|AED|SAR)\s*[\d][\d\s,\.]*)",
            r"((?:\$|€|£|₹|₦|¥)\s*[\d][\d\s,\.]*)",
        ]

        for pattern in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                fragment = match.group(1)
                raw_value = re.sub(r"^(?:USD|EUR|GBP|CAD|AUD|INR|NGN|KES|JPY|CNY|CHF|AED|SAR|\$|€|£|₹|₦|¥)\s*", "", fragment, flags=re.IGNORECASE)
                value = self._parse_amount_value(raw_value, source_fragment=fragment)
                if value is None or value < 0:
                    continue
                context = text[max(0, match.start() - 40):match.end() + 40].lower()
                candidates.append(
                    {
                        "value": value,
                        "raw": raw_value,
                        "currency": self._detect_currency(f"{fragment} {context}"),
                        "score": 1 + self._score_amount_candidate(fragment, context),
                    }
                )

        if not candidates:
            return []

        unique_map: Dict[float, Dict[str, Any]] = {}
        for candidate in candidates:
            value = candidate["value"]
            current = unique_map.get(value)
            if not current or candidate.get("score", 0) > current.get("score", 0):
                unique_map[value] = candidate

        return sorted(unique_map.values(), key=lambda x: (x.get("score", 0), x["value"]), reverse=True)

    def _looks_like_identifier_token(self, raw: str) -> bool:
        """Reject ID-like numeric tokens that should not be treated as monetary values."""
        token = re.sub(r'[^A-Za-z0-9.-]', '', str(raw or ""))
        if not token:
            return False
        # Long digit strings are usually invoice/transaction identifiers.
        if re.fullmatch(r"\d{8,}", token):
            return True
        # Mixed ID forms such as INV-12345 should not flow through amount parsing.
        if re.search(r"[A-Za-z]", token) and re.search(r"\d", token):
            return True
        return False

    def _filter_amounts_against_invoice_numbers(
        self,
        amounts: List[Dict[str, Any]],
        invoice_numbers: List[str]
    ) -> List[Dict[str, Any]]:
        invoice_digits = {
            re.sub(r'\D', '', str(number))
            for number in invoice_numbers
            if number
        }
        if not invoice_digits:
            return amounts

        filtered = []
        for amount in amounts:
            raw = amount.get('raw') if isinstance(amount, dict) else amount
            digits = re.sub(r'\D', '', str(raw))
            trimmed = digits.rstrip('0') if digits else digits
            if digits:
                if digits in invoice_digits or trimmed in invoice_digits:
                    continue
                # If the amount digits contain the invoice digits (or vice versa), skip.
                if any(inv in digits or digits in inv for inv in invoice_digits):
                    continue
            filtered.append(amount)
        return filtered
    
    def _extract_invoice_numbers(self, text: str) -> List[str]:
        """Extract invoice numbers from text."""
        candidates: List[Dict[str, Any]] = []

        for pattern in self.INVOICE_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                candidate = self._normalize_invoice_candidate(match.group(1))
                if not candidate:
                    continue
                context = text[max(0, match.start() - 30):match.end() + 30].lower()
                label = match.group(0).lower()
                score = 0
                if "invoice" in label or label.startswith("inv"):
                    score += 4
                elif "bill" in label or "reference" in label:
                    score += 2
                elif "order" in label or "po" in label or "receipt" in label:
                    score += 1
                if re.search(r"(invoice|bill|reference|payment)", context):
                    score += 1
                if re.search(r"[A-Za-z]", candidate) and re.search(r"\d", candidate):
                    score += 1
                if len(candidate) > 28:
                    score -= 1
                candidates.append({"value": candidate, "score": score, "start": match.start()})

        # Fallback for common free-form invoice IDs in subject/body.
        fallback_patterns = [
            r"\b(?:invoice|inv|bill|doc)\s*[:#-]?\s*([A-Z0-9][A-Z0-9/_-]{3,})\b",
            r"\b([A-Z]{1,4}-\d{4,})\b",
        ]
        for pattern in fallback_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                candidate = self._normalize_invoice_candidate(match.group(1))
                if candidate:
                    candidates.append({"value": candidate, "score": 1, "start": match.start()})

        if not candidates:
            return []

        # Keep highest score per normalized value.
        best_by_value: Dict[str, Dict[str, Any]] = {}
        for candidate in candidates:
            key = candidate["value"].upper()
            current = best_by_value.get(key)
            if not current or candidate["score"] > current["score"]:
                best_by_value[key] = candidate

        ranked = sorted(best_by_value.values(), key=lambda item: (-item["score"], item["start"]))
        return [item["value"] for item in ranked]

    def _normalize_invoice_candidate(self, value: Optional[str]) -> Optional[str]:
        """Normalize/validate invoice candidate tokens."""
        token = str(value or "").strip().strip(".,;:()[]{}")
        token = token.replace(" ", "")
        if not token:
            return None
        if not self._is_probable_invoice_number(token):
            return None
        return token

    def _is_probable_invoice_number(self, token: str) -> bool:
        """Heuristics to keep valid invoice IDs and drop dates/noise."""
        if len(token) < 4 or len(token) > 40:
            return False
        if not re.search(r"\d", token):
            return False
        if re.fullmatch(r"\d{4}", token) and 1900 <= int(token) <= 2100:
            return False
        if re.fullmatch(r"\d+\.\d{2}", token):
            return False
        if re.fullmatch(r"\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}", token):
            return False
        # Pure short numerics are typically line references rather than invoice IDs.
        if re.fullmatch(r"\d{1,5}", token):
            return False
        return True
    
    def _extract_dates(self, text: str) -> List[str]:
        """Extract and validate dates from text."""
        dates = []
        
        for pattern in self.DATE_PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        
        # Normalize to ISO format with validation
        normalized = []
        date_formats = [
            '%Y-%m-%d',
            '%Y/%m/%d',
            '%d/%m/%Y',
            '%m/%d/%Y',
            '%d-%m-%Y',
            '%m-%d-%Y',
            '%d.%m.%Y',
            '%d %B %Y',
            '%d %b %Y',
            '%B %d, %Y',
            '%B %d %Y',
            '%b %d, %Y',
            '%b %d %Y',
        ]
        
        for d in dates:
            d = d.strip()
            parsed_date = None
            
            for fmt in date_formats:
                try:
                    parsed_date = datetime.strptime(d, fmt)
                    break
                except ValueError:
                    continue
            
            if parsed_date:
                # Validate date is reasonable (not too far in past or future)
                if self._validate_date(parsed_date):
                    normalized.append(parsed_date.strftime('%Y-%m-%d'))
        
        ordered_unique: List[str] = []
        seen = set()
        for date_value in normalized:
            if date_value in seen:
                continue
            seen.add(date_value)
            ordered_unique.append(date_value)

        return ordered_unique
    
    def _validate_date(self, date: datetime) -> bool:
        """Validate that a date is reasonable for a financial document."""
        now = datetime.now(timezone.utc)
        
        # Date shouldn't be more than 2 years in the past
        min_date = datetime(now.year - 2, 1, 1)
        
        # Date shouldn't be more than 1 year in the future
        max_date = datetime(now.year + 1, 12, 31)
        
        return min_date <= date <= max_date
    
    def _extract_due_date(self, text: str) -> Optional[str]:
        """Extract due date specifically."""
        patterns = [
            r'Due\s*(?:Date)?[:\s]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
            r'Payment\s+Due[:\s]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
            r'Due\s+by[:\s]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                dates = self._extract_dates(match.group(1))
                if dates:
                    return dates[0]
        
        return None
    
    def _extract_line_items(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract line items from invoice text.
        Handles both structured tables and free-form text.
        """
        items = []
        
        # Check if text contains table markers (from pdfplumber extraction)
        if '|' in text:
            items.extend(self._extract_line_items_from_table(text))
        
        # Also try regex-based extraction
        items.extend(self._extract_line_items_regex(text))
        
        # Deduplicate by description similarity
        unique_items = []
        seen_descs = set()
        for item in items:
            desc_lower = item['description'].lower()[:30]
            if desc_lower not in seen_descs:
                seen_descs.add(desc_lower)
                unique_items.append(item)
        
        return unique_items[:30]  # Limit to 30 items
    
    def _extract_line_items_from_table(self, text: str) -> List[Dict[str, Any]]:
        """Extract line items from table-formatted text."""
        items = []
        
        # Process lines that contain pipe separators (table rows)
        for line in text.split('\n'):
            if '|' not in line:
                continue
            
            cells = [c.strip() for c in line.split('|')]
            
            # Skip header rows
            if any(h in line.lower() for h in ['description', 'item', 'qty', 'quantity', 'price', 'amount', 'total']):
                continue
            
            # Try to identify description and amount columns
            description = None
            amount = None
            quantity = None
            unit_price = None
            
            for cell in cells:
                if not cell:
                    continue
                
                # Check if cell is a number/amount
                amount_match = re.search(r'^[\$€£₦]?\s*([\d,]+\.?\d*)\s*$', cell)
                if amount_match:
                    val = float(amount_match.group(1).replace(',', ''))
                    if val > 0:
                        if unit_price is None and val < 10000:
                            unit_price = val
                        elif amount is None:
                            amount = val
                        elif val > amount:
                            unit_price = amount
                            amount = val
                elif len(cell) > 3 and not cell.isdigit():
                    # Likely a description
                    if description is None or len(cell) > len(description):
                        description = cell
            
            if description and amount and amount > 0:
                item = {
                    "description": description[:100],
                    "amount": amount
                }
                if quantity:
                    item["quantity"] = quantity
                if unit_price:
                    item["unit_price"] = unit_price
                items.append(item)
        
        return items
    
    def _extract_line_items_regex(self, text: str) -> List[Dict[str, Any]]:
        """Extract line items using regex patterns."""
        items = []
        
        # Pattern for line items: description followed by amount
        patterns = [
            # Description ... Amount
            r'^(.{10,60}?)\s+([\d,]+\.\d{2})\s*$',
            # Quantity x Description @ Price = Amount
            r'^(\d+)\s*[xX×]\s*(.{5,50}?)\s*[@at]\s*[\$€£]?\s*([\d,]+\.?\d*)\s*=?\s*[\$€£]?\s*([\d,]+\.?\d*)',
            # Description (Price)
            r'^([A-Z][^0-9]{5,40})\s+[\$€£]?\s*([\d,]+\.\d{2})',
        ]
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:
                continue
            
            # Skip obvious non-item lines
            if any(kw in line.lower() for kw in ['total', 'subtotal', 'tax', 'shipping', 'discount', 'invoice', 'date', 'due']):
                continue
            
            for pattern in patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    groups = match.groups()
                    if len(groups) >= 2:
                        desc = groups[0].strip() if isinstance(groups[0], str) else str(groups[0])
                        amount_str = groups[-1] if len(groups[-1]) > 0 else groups[-2]
                        
                        try:
                            amount = float(str(amount_str).replace(',', ''))
                            if amount > 0 and len(desc) > 3:
                                items.append({
                                    "description": desc[:100],
                                    "amount": amount
                                })
                                break
                        except (ValueError, TypeError):
                            continue
        
        return items
    
    def _extract_party(self, text: str, direction: str) -> Optional[str]:
        """Extract payer or payee from text."""
        patterns = [
            rf'{direction}[:\s]+([A-Za-z\s]+)',
            rf'(?:Paid|Payment)\s+{direction}[:\s]+([A-Za-z\s]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _detect_currency(self, text: str) -> str:
        """Detect currency from text with comprehensive symbol/code support."""
        text_lower = text.lower()

        # Avoid false positives such as "... for 1 Dec 2025 ...".
        if re.search(r"\bR\s*\d", text):
            return "ZAR"
        
        # Check for currency symbols and codes
        currency_indicators = [
            ('€', 'EUR'),
            ('eur', 'EUR'),
            ('£', 'GBP'),
            ('gbp', 'GBP'),
            ('₦', 'NGN'),
            ('ngn', 'NGN'),
            ('naira', 'NGN'),
            ('gh₵', 'GHS'),
            ('ghs', 'GHS'),
            ('ghana cedi', 'GHS'),
            ('cedi', 'GHS'),
            ('zar', 'ZAR'),
            ('rand', 'ZAR'),
            ('kes', 'KES'),
            ('ksh', 'KES'),
            ('shilling', 'KES'),
            ('¥', 'JPY'),
            ('jpy', 'JPY'),
            ('yen', 'JPY'),
            ('cny', 'CNY'),
            ('rmb', 'CNY'),
            ('yuan', 'CNY'),
            ('₹', 'INR'),
            ('inr', 'INR'),
            ('rupee', 'INR'),
            ('rs.', 'INR'),
            ('chf', 'CHF'),
            ('franc', 'CHF'),
            ('a$', 'AUD'),
            ('aud', 'AUD'),
            ('c$', 'CAD'),
            ('cad', 'CAD'),
            ('sek', 'SEK'),
            ('kr', 'SEK'),  # Could be SEK/NOK/DKK
            ('nok', 'NOK'),
            ('dkk', 'DKK'),
            ('pln', 'PLN'),
            ('zł', 'PLN'),
            ('zloty', 'PLN'),
            ('r$', 'BRL'),
            ('brl', 'BRL'),
            ('real', 'BRL'),
            ('mxn', 'MXN'),
            ('peso', 'MXN'),
            ('aed', 'AED'),
            ('dirham', 'AED'),
            ('sar', 'SAR'),
            ('riyal', 'SAR'),
            ('sgd', 'SGD'),
            ('hkd', 'HKD'),
            ('hk$', 'HKD'),
            ('nzd', 'NZD'),
            ('nz$', 'NZD'),
            ('thb', 'THB'),
            ('baht', 'THB'),
            ('฿', 'THB'),
            ('$', 'USD'),
            ('usd', 'USD'),
            ('dollar', 'USD'),
        ]
        
        for indicator, currency in currency_indicators:
            if indicator in text_lower:
                return currency
        
        return 'USD'  # Default to USD as most common

    # Zip-bomb guardrails. A hostile sender can ship a 100KB ZIP that
    # expands to many GB because `info.file_size` is read from the
    # archive's own central directory — i.e. attacker-controlled. We
    # bound two dimensions:
    #   - per-file expansion ratio (file_size / compress_size)
    #   - total decompressed bytes across the whole archive
    # Anything that trips either limit is skipped (we don't raise — the
    # caller already treats an empty extraction list as "no usable
    # attachments" and falls back to the email body).
    _MAX_ZIP_RATIO = 100
    _MAX_ZIP_TOTAL_DECOMPRESSED = 200 * 1024 * 1024  # 200 MB across all members

    def _extract_archive_attachments(self, content_base64: str, filename: str) -> List[Dict[str, Any]]:
        """Extract files from a ZIP archive attachment."""
        try:
            data = base64.b64decode(content_base64)
            extracted: List[Dict[str, Any]] = []
            total_decompressed = 0
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for info in zf.infolist():
                    if info.is_dir():
                        continue
                    # Only process common invoice file types
                    name_lower = info.filename.lower()
                    if not any(name_lower.endswith(ext) for ext in ('.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.tif')):
                        continue
                    # Size limit: skip files > 25MB
                    if info.file_size > 25 * 1024 * 1024:
                        continue
                    # Zip-bomb checks — attacker controls the claimed
                    # file_size, but any extreme compression ratio is a
                    # red flag regardless of what the header claims.
                    if info.compress_size > 0 and info.file_size / info.compress_size > self._MAX_ZIP_RATIO:
                        logger.warning(
                            "Skipping suspicious archive member %s: ratio %d exceeds %d",
                            info.filename,
                            int(info.file_size / max(info.compress_size, 1)),
                            self._MAX_ZIP_RATIO,
                        )
                        continue
                    if total_decompressed + info.file_size > self._MAX_ZIP_TOTAL_DECOMPRESSED:
                        logger.warning(
                            "Stopping archive %s extraction: total decompressed size would exceed %d bytes",
                            filename,
                            self._MAX_ZIP_TOTAL_DECOMPRESSED,
                        )
                        break
                    file_data = zf.read(info.filename)
                    total_decompressed += len(file_data)
                    # Determine content type
                    if name_lower.endswith('.pdf'):
                        content_type = "application/pdf"
                    elif name_lower.endswith('.png'):
                        content_type = "image/png"
                    elif name_lower.endswith(('.tiff', '.tif')):
                        content_type = "image/tiff"
                    else:
                        content_type = "image/jpeg"
                    extracted.append({
                        "filename": info.filename,
                        "name": info.filename,
                        "content_base64": base64.b64encode(file_data).decode("utf-8"),
                        "content_type": content_type,
                        "size": info.file_size,
                        "source": f"archive:{filename}",
                    })
            return extracted
        except Exception as exc:
            logger.warning("Failed to extract archive %s: %s", filename, exc)
            return []

    def _parse_attachment(self, attachment: Dict) -> Optional[Dict[str, Any]]:
        """Parse an email attachment."""
        name = (attachment.get('name') or attachment.get('filename') or '').lower()
        content_type = attachment.get('content_type') or attachment.get('mime_type') or ''
        content_base64 = attachment.get('content_base64')
        content_text = attachment.get('content_text')
        
        # Determine attachment type
        if 'pdf' in content_type or name.endswith('.pdf'):
            parsed_text = None
            if content_text:
                parsed_text = content_text
            elif content_base64:
                pdf_result = self._extract_pdf_text(content_base64)
                # C8: _extract_pdf_text may return a dict on base64 corruption
                if isinstance(pdf_result, dict) and pdf_result.get("status") == "attachment_corrupted":
                    return {
                        "name": attachment.get('name') or attachment.get('filename'),
                        "type": "document",
                        "content_type": "application/pdf",
                        "requires_ocr": False,
                        "parsed": False,
                        "status": "attachment_corrupted",
                        "error": pdf_result.get("error"),
                    }
                # C11: _extract_pdf_text may return a dict for password-protected PDFs
                if isinstance(pdf_result, dict) and pdf_result.get("status") == "attachment_password_protected":
                    return {
                        "name": attachment.get('name') or attachment.get('filename'),
                        "type": "document",
                        "content_type": "application/pdf",
                        "requires_ocr": False,
                        "parsed": False,
                        "status": "attachment_password_protected",
                    }
                parsed_text = pdf_result

            parsed_invoice = None
            if parsed_text:
                parsed_invoice = self.parse_invoice_text(parsed_text)
            parsed_type = (
                str(parsed_invoice.get("type") or "").strip().lower()
                if isinstance(parsed_invoice, dict)
                else ""
            )

            ocr_status = None
            if not parsed_text and not (OCR_AVAILABLE and PDFIUM_AVAILABLE):
                ocr_status = "requires_ocr"
                logger.warning(
                    "PDF attachment %r has no extractable text and OCR is unavailable",
                    attachment.get('name') or attachment.get('filename'),
                )

            return {
                "name": attachment.get('name') or attachment.get('filename'),
                "type": parsed_type or ("invoice" if 'invoice' in name else "document"),
                "content_type": "application/pdf",
                "requires_ocr": False if parsed_text else True,
                "ocr_status": ocr_status,
                "parsed": bool(parsed_text),
                "content_text": parsed_text,
                "extraction": parsed_invoice
            }
        elif 'word' in content_type or name.endswith('.docx'):
            parsed_text = None
            if content_text:
                parsed_text = content_text
            elif content_base64:
                parsed_text = self._extract_docx_text(content_base64)

            parsed_invoice = None
            if parsed_text:
                parsed_invoice = self.parse_invoice_text(parsed_text)
            parsed_type = (
                str(parsed_invoice.get("type") or "").strip().lower()
                if isinstance(parsed_invoice, dict)
                else ""
            )

            return {
                "name": attachment.get('name') or attachment.get('filename'),
                "type": parsed_type or ("invoice" if 'invoice' in name else "document"),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "requires_ocr": False,
                "parsed": bool(parsed_text),
                "content_text": parsed_text,
                "extraction": parsed_invoice
            }
        elif 'csv' in content_type or name.endswith('.csv'):
            return {
                "name": attachment.get('name') or attachment.get('filename'),
                "type": "statement" if 'statement' in name else "data",
                "content_type": "text/csv",
                "requires_ocr": False,
                "parsed": False
            }
        elif 'excel' in content_type or name.endswith(('.xlsx', '.xls')):
            return {
                "name": attachment.get('name') or attachment.get('filename'),
                "type": "spreadsheet",
                "content_type": "application/excel",
                "requires_ocr": False,
                "parsed": False
            }
        elif 'image' in content_type or name.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.tiff', '.bmp')):
            # Attempt OCR extraction
            ocr_text = None
            parsed_invoice = None
            ocr_status = None

            if content_base64 and OCR_AVAILABLE:
                ocr_text = self._extract_image_text_ocr(content_base64)
                if ocr_text:
                    parsed_invoice = self.parse_invoice_text(ocr_text)
            elif content_base64 and not OCR_AVAILABLE:
                ocr_status = "requires_ocr"
                logger.warning(
                    "Scanned image attachment %r cannot be processed - OCR (pytesseract) is not installed",
                    attachment.get('name') or attachment.get('filename'),
                )

            return {
                "name": attachment.get('name') or attachment.get('filename'),
                "type": "invoice" if 'invoice' in name else "document",
                "content_type": content_type,
                "requires_ocr": not bool(ocr_text),
                "ocr_status": ocr_status,
                "parsed": bool(ocr_text),
                "content_text": ocr_text,
                "extraction": parsed_invoice
            }

        return None
    
    def _extract_image_text_ocr(self, content_base64: str) -> Optional[str]:
        """
        Extract text from an image using OCR (pytesseract).
        
        Args:
            content_base64: Base64-encoded image content
            
        Returns:
            Extracted text or None if OCR fails
        """
        if not OCR_AVAILABLE:
            logger.warning("OCR not available - pytesseract not installed")
            return None
        
        try:
            # Decode base64 image
            image_data = base64.b64decode(content_base64)
            image = Image.open(io.BytesIO(image_data))
            return self._extract_pil_text_ocr(image)
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
            return None

    def _extract_pil_text_ocr(self, image: "Image.Image") -> Optional[str]:
        """Extract OCR text from a PIL image after invoice-friendly preprocessing."""
        if not OCR_AVAILABLE:
            return None

        try:
            # Convert to RGB if necessary (for PNG with transparency / palettes)
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if 'A' in image.mode else None)
                image = background

            if image.mode != 'L':
                gray = image.convert('L')
            else:
                gray = image

            from PIL import ImageEnhance

            contrast = ImageEnhance.Contrast(gray).enhance(1.8)
            sharpened = ImageEnhance.Sharpness(contrast).enhance(1.3)

            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(sharpened, config=custom_config)

            if text and len(text.strip()) > 20:
                logger.info("OCR extracted %s characters from rendered image", len(text))
                return text.strip()
            return None
        except Exception as e:
            logger.warning(f"PIL OCR extraction failed: {e}")
            return None

    def _extract_docx_text(self, content_base64: str) -> Optional[str]:
        """
        Extract text from a DOCX attachment.
        """
        if not DOCX_AVAILABLE:
            return None

        try:
            raw = base64.b64decode(content_base64)
            document = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in document.paragraphs if p.text]
            text = "\n".join(paragraphs).strip()
            return text or None
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            return None
    
    def extract_from_image(self, image_path: str) -> Optional[Dict[str, Any]]:
        """
        Extract invoice data from an image file.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Parsed invoice data or None
        """
        if not OCR_AVAILABLE:
            return None
        
        try:
            with open(image_path, 'rb') as f:
                content = base64.b64encode(f.read()).decode()
            
            text = self._extract_image_text_ocr(content)
            if text:
                return self.parse_invoice_text(text)
            
            return None
        except Exception as e:
            logger.warning(f"Failed to extract from image {image_path}: {e}")
            return None

    def _extract_pdf_text(self, content_base64: str, max_pages: int = None):
        """
        Extract text from a base64-encoded PDF attachment.
        Uses text-layer extraction first and falls back to OCR for scanned PDFs.

        Args:
            content_base64: Base64-encoded PDF content
            max_pages: Maximum pages to process (None = all pages)

        Returns:
            Extracted text string, or a dict with status/error if base64 is corrupted.
        """
        try:
            data = base64.b64decode(content_base64)
        except Exception as e:
            logger.warning(f"Failed to decode PDF base64: {e}")
            # C8: Return structured error so callers can detect corrupted attachments
            return {"status": "attachment_corrupted", "error": str(e)}
        return self._extract_pdf_text_from_bytes(data, max_pages=max_pages)

    def _extract_pdf_text_from_bytes(self, pdf_data: bytes, max_pages: int = None):
        """Extract PDF text using text-layer parsing first, then OCR when needed."""
        text_candidates: List[Tuple[str, str]] = []

        text_layer = self._extract_pdf_text_layer(pdf_data, max_pages=max_pages)
        # C11: Propagate password-protected dict up the call chain
        if isinstance(text_layer, dict):
            return text_layer
        parsed_text_layer = self.parse_invoice_text(text_layer) if text_layer else None
        if text_layer:
            text_candidates.append(("text_layer", text_layer))

        if self._should_attempt_pdf_ocr(text_layer, parsed_text_layer):
            ocr_text = self._extract_pdf_text_ocr(pdf_data, max_pages=max_pages)
            if ocr_text:
                text_candidates.append(("ocr", ocr_text))

        best = self._choose_best_pdf_text_candidate(text_candidates)
        return best[1] if best else None

    def _extract_pdf_text_layer(self, pdf_data: bytes, max_pages: int = None):
        """Extract PDF text from embedded text layers before attempting OCR."""
        if PDFPLUMBER_AVAILABLE:
            try:
                text = self._extract_with_pdfplumber(pdf_data, max_pages)
                if text:
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")

        # C11: _extract_with_pypdf2 may return a dict for password-protected PDFs;
        # propagate it up so callers can detect the condition.
        return self._extract_with_pypdf2(pdf_data, max_pages=max_pages)

    def _extract_with_pypdf2(self, pdf_data: bytes, max_pages: int = None):
        """Extract PDF text using PyPDF2 as a fallback text-layer parser.

        Returns:
            Extracted text string, or a dict with status if password-protected, or None.
        """
        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))

            # C11: Detect password-protected PDFs before extraction
            if reader.is_encrypted:
                try:
                    # Attempt empty password (some PDFs are "encrypted" with no password)
                    if not reader.decrypt(""):
                        logger.warning("PDF is password-protected and cannot be decrypted")
                        return {"status": "attachment_password_protected"}
                except Exception:
                    logger.warning("PDF is password-protected and cannot be decrypted")
                    return {"status": "attachment_password_protected"}

            total_pages = len(reader.pages)
            pages_to_read = total_pages if max_pages is None else min(total_pages, max_pages)

            text_parts = []
            for i in range(pages_to_read):
                page = reader.pages[i]
                extracted = page.extract_text() or ""
                text_parts.append(extracted)

            text = "\n".join(text_parts).strip()
            return text or None
        except Exception as e:
            # C11: Catch specific password-related errors from PyPDF2
            error_str = str(e).lower()
            if "password" in error_str or "encrypted" in error_str:
                logger.warning("PDF is password-protected: %s", e)
                return {"status": "attachment_password_protected"}
            logger.warning(f"PyPDF2 extraction failed: {e}")
            return None
    
    def _extract_with_pdfplumber(self, pdf_data: bytes, max_pages: int = None) -> Optional[str]:
        """Extract text and tables from PDF using pdfplumber."""
        if not PDFPLUMBER_AVAILABLE:
            return None
        
        try:
            with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
                total_pages = len(pdf.pages)
                pages_to_read = total_pages if max_pages is None else min(total_pages, max_pages)
                
                all_text = []
                
                for i in range(pages_to_read):
                    page = pdf.pages[i]
                    
                    # Extract tables first
                    tables = page.extract_tables()
                    table_text = []
                    for table in tables:
                        if table:
                            for row in table:
                                if row:
                                    row_text = " | ".join(str(cell or '') for cell in row)
                                    table_text.append(row_text)
                    
                    # Extract regular text
                    page_text = page.extract_text() or ""
                    
                    # Combine table text and page text
                    if table_text:
                        all_text.append(f"--- Page {i+1} Tables ---")
                        all_text.extend(table_text)
                    all_text.append(f"--- Page {i+1} Text ---")
                    all_text.append(page_text)
                
                return "\n".join(all_text).strip() or None
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            return None

    def _extract_pdf_text_ocr(self, pdf_data: bytes, max_pages: int = None) -> Optional[str]:
        """Render a PDF to images and OCR the pages when no reliable text layer exists."""
        if not (OCR_AVAILABLE and PDFIUM_AVAILABLE):
            return None

        try:
            with pdfium.PdfDocument(pdf_data) as doc:
                total_pages = len(doc)
                page_limit = max_pages if max_pages is not None else 3
                pages_to_read = min(total_pages, max(1, page_limit))
                scale = 300 / 72.0
                text_parts: List[str] = []

                for index in range(pages_to_read):
                    page = doc[index]
                    bitmap = None
                    try:
                        bitmap = page.render(scale=scale, grayscale=True)
                        image = bitmap.to_pil()
                        page_text = self._extract_pil_text_ocr(image)
                        if page_text:
                            text_parts.append(f"--- Page {index + 1} OCR ---")
                            text_parts.append(page_text)
                    finally:
                        if bitmap is not None and hasattr(bitmap, "close"):
                            bitmap.close()
                        if hasattr(page, "close"):
                            page.close()

            text = "\n".join(text_parts).strip()
            return text or None
        except Exception as e:
            logger.warning(f"PDF OCR extraction failed: {e}")
            return None

    def _should_attempt_pdf_ocr(
        self,
        text: Optional[str],
        parsed_invoice: Optional[Dict[str, Any]],
    ) -> bool:
        """Decide whether a PDF should be rasterized and OCR'd."""
        if not (OCR_AVAILABLE and PDFIUM_AVAILABLE):
            # Log when a scanned PDF would need OCR but we can't provide it
            normalized = re.sub(r"\s+", "", str(text or ""))
            if len(normalized) < 60:
                logger.warning(
                    "Scanned PDF detected (text layer too thin: %d chars) but OCR is unavailable. "
                    "Install pytesseract + pypdfium2 to process scanned invoices.",
                    len(normalized),
                )
            return False

        normalized = re.sub(r"\s+", "", str(text or ""))
        if len(normalized) < 60:
            return True

        if not isinstance(parsed_invoice, dict):
            return True

        strong_fields = 0
        if str(parsed_invoice.get("vendor") or "").strip():
            strong_fields += 1
        if str(parsed_invoice.get("invoice_number") or "").strip():
            strong_fields += 1

        amount = parsed_invoice.get("amount")
        amount_value = None
        if isinstance(amount, dict):
            amount_value = amount.get("value")
        elif amount is not None:
            amount_value = amount
        if amount_value is not None:
            strong_fields += 1

        if str(parsed_invoice.get("date") or "").strip():
            strong_fields += 1
        if str(parsed_invoice.get("due_date") or "").strip():
            strong_fields += 1

        text_lower = str(text or "").lower()
        has_table_noise = text_lower.count("--- page") > 0 and text_lower.count(" | ") > 4
        if strong_fields >= 4 and not has_table_noise:
            return False
        return strong_fields <= 2 or has_table_noise

    def _choose_best_pdf_text_candidate(
        self,
        candidates: List[Tuple[str, str]],
    ) -> Optional[Tuple[str, str]]:
        """Choose the strongest PDF text candidate by parsed invoice signal."""
        best: Optional[Tuple[str, str]] = None
        best_score = -1
        best_text_length = -1

        for method, text in candidates:
            parsed = self.parse_invoice_text(text) if text else {}
            score = self._score_invoice_parse(parsed)
            text_length = len(re.sub(r"\s+", "", text or ""))
            if score > best_score or (score == best_score and text_length > best_text_length):
                best = (method, text)
                best_score = score
                best_text_length = text_length

        return best

    def _score_invoice_parse(self, parsed_invoice: Optional[Dict[str, Any]]) -> int:
        """Score parsed invoice completeness for source arbitration."""
        if not isinstance(parsed_invoice, dict):
            return 0

        score = 0
        vendor = str(parsed_invoice.get("vendor") or "").strip()
        if vendor and not self._is_vendor_noise_candidate(vendor):
            score += 2

        if str(parsed_invoice.get("invoice_number") or "").strip():
            score += 3

        amount = parsed_invoice.get("amount")
        amount_value = None
        if isinstance(amount, dict):
            amount_value = amount.get("value")
        elif amount is not None:
            amount_value = amount
        if amount_value is not None:
            score += 3 if float(amount_value) > 0 else 2

        if str(parsed_invoice.get("date") or "").strip():
            score += 1
        if str(parsed_invoice.get("due_date") or "").strip():
            score += 1

        line_items = parsed_invoice.get("line_items")
        if isinstance(line_items, list) and line_items:
            score += 1

        return score
    
    def _calculate_confidence(
        self,
        email_type: str,
        amounts: List[Dict],
        invoice_numbers: List[str]
    ) -> float:
        """Calculate confidence score for parsed data."""
        score = 0.0
        
        # Base score for AP email type
        if email_type in ['invoice', 'payment_request', 'refund', 'credit_note']:
            score += 0.3
        
        # Score for extracted amounts
        if amounts:
            score += 0.3
            if len(amounts) == 1:  # Single clear amount
                score += 0.1
        
        # Score for invoice numbers
        if invoice_numbers:
            score += 0.2
            if len(invoice_numbers) == 1:  # Single clear invoice
                score += 0.1
        
        return min(score, 1.0)


# Convenience functions

def parse_email(
    subject: str,
    body: str,
    sender: str,
    attachments: List[Dict] = None,
    *,
    organization_id: str,
    thread_id: str = None,
) -> Dict[str, Any]:
    """Parse an email and extract financial data.

    Primary path: LLMEmailParser (Claude Haiku for text, Sonnet for vision).
    Automatic fallback: regex EmailParser when Claude is unavailable or fails.

    organization_id is required — vendor history / past corrections /
    thread context all key off it, and a missing org would silently
    bind extracted invoices to the legacy ``default`` tenant.
    """
    organization_id = assert_org_id(organization_id, context="parse_email")
    from clearledgr.services.llm_email_parser import parse_email_with_llm
    return parse_email_with_llm(
        subject, body, sender, attachments,
        organization_id=organization_id,
        thread_id=thread_id,
    )


def parse_invoice_text(text: str) -> Dict[str, Any]:
    """Parse invoice text."""
    parser = EmailParser()
    return parser.parse_invoice_text(text)


def parse_payment_confirmation(text: str) -> Dict[str, Any]:
    """Parse payment confirmation."""
    parser = EmailParser()
    return parser.parse_payment_confirmation(text)


def get_parser_capabilities() -> Dict[str, Any]:
    """
    Get the current capabilities of the email parser.
    Useful for checking what features are available.
    """
    return {
        "ocr_available": OCR_AVAILABLE,
        "table_extraction_available": PDFPLUMBER_AVAILABLE,
        "pdf_ocr_available": OCR_AVAILABLE and PDFIUM_AVAILABLE,
        "fuzzy_matching_available": FUZZY_AVAILABLE,
        "supported_image_formats": [".png", ".jpg", ".jpeg", ".gif", ".webp", ".tiff", ".bmp"] if OCR_AVAILABLE else [],
        "supported_document_formats": [".pdf", ".csv", ".xlsx", ".xls"],
        "supported_currencies": EmailParser().supported_currencies,
        "known_vendors_count": len(KNOWN_VENDORS),
        "recommendations": _get_recommendations()
    }


def _get_recommendations() -> List[str]:
    """Get recommendations for improving parser capabilities."""
    recommendations = []
    
    if not OCR_AVAILABLE:
        recommendations.append("Install pytesseract and pillow for OCR support: pip install pytesseract pillow")
    
    if not PDFPLUMBER_AVAILABLE:
        recommendations.append("Install pdfplumber for better PDF table extraction: pip install pdfplumber")
    
    if not FUZZY_AVAILABLE:
        recommendations.append("Install rapidfuzz for fuzzy vendor matching: pip install rapidfuzz")
    
    return recommendations


def extract_from_image_file(image_path: str) -> Optional[Dict[str, Any]]:
    """
    Extract invoice data from an image file using OCR.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Parsed invoice data or None if OCR is not available
    """
    parser = EmailParser()
    return parser.extract_from_image(image_path)
